from docx import Document
import re
import csv
from time import sleep
def shuffle_checker(iterable):
    if 'كل ما سبق صحيح' in iterable or '(1+2)' in iterable or '(2+1)' in iterable or \
            '(2+3)' in iterable or '(1+3)' in iterable or 'كل ما سبق خاطئ' in iterable \
            or 'غير ذلك' in iterable or '(3+1)' in iterable or '(3+2)' in iterable or '(2 + 1)' in iterable or \
            '(3 + 1)' in iterable or '(3 + 2)' in iterable or '(1 + 2)' in iterable or '(2 + 3)' in iterable or \
            '(3 + 2)' in iterable or '(2 + 3)' in iterable or 'ليس أيّاً مما سبق' in iterable or 'كل ما سبق' in iterable:
        return True
    else:
        return False
checker = True
while checker:
    file = input('Please enter the file\'s directory:')
    output = input('Please enter the directory for the new file with the desired name: ')
    style = eval(input('''style 1 is numbered questions and colored correct answers
style 2 is numbered answers and an asterisk 
Please enter the style: '''))
    if output.split('.')[-1] != 'csv':
        print('The output file is not in csv format')
        checker = True
    elif file.split('.')[-1] != 'docx':
        print('The input file is not in word format')
        checker = True
    else:
        checker = False
doc = Document(file)
questions = []
correct_answers = []
questions_and_answers = []  # List to store questions and answers
current_question = None
current_answers = []
target_hex_color = 'FF0000'
target_rgb = tuple(int(target_hex_color[i:i + 2], 16) for i in (0, 2, 4))
if style == 1:
    for paragraph in doc.paragraphs:
        line = paragraph.text.strip().split()
        if len(line) == 0:
            continue
        if line[0][0].isdigit():
            line[0] = line[0].replace('-', '')
            line[0] = line[0].replace('ـ', '')
            line[0] = int(line[0])
            questions.append(line)
        line_text = paragraph.text.strip()

        if line_text and line_text[0].isdigit():
            if current_question:
                questions_and_answers.append({
                    "question": current_question,
                    "answers": current_answers
                })

                # Start a new question
            current_question = line_text
            current_answers = []

            # Otherwise, consider it as an answer and add it to the current_answers list
        elif current_question:
            current_answers.append(line_text)

        for run in paragraph.runs:
            if run.font.color.rgb == target_rgb:
                if len(run.text.strip()) <= 1:
                    continue
                else:
                    correct_answers.append(run.text.strip())
    acc = 0
    if current_question:
        questions_and_answers.append({
            "question": current_question,
            "answers": current_answers
        })
    for i in questions:
        acc += 1
        if int(i[0]) != acc:
            print(f'near {int(i[0])} there is a problem in the number of questions')
            acc += 1

    counter = 1
    for qa in questions_and_answers:
        input_string = qa['answers']
        if counter == len(questions_and_answers):
            input_string = input_string[0]
        else:
            if len(input_string) == 1:
                input_string = input_string[0]
            elif len(input_string) == 2:
                input_string = input_string[0] + '    ' + input_string[1]
            elif len(input_string) == 3:
                input_string = input_string[0] + '    ' + input_string[1] + '     ' + input_string[2]
        split_strings = re.split(r'\s\s', input_string)
        # Filter out any empty strings resulting from consecutive spaces
        split_strings = [s.strip() for s in split_strings if s]
        qa['answers'] = split_strings
        counter += 1
    for qa in questions_and_answers:
        qa['question'] = re.sub('\d+.\s', "", qa['question']).strip()

    question_counter = 0
    with open(output, mode='w', encoding="utf-8-sig", newline='') as file:
        writer = csv.writer(file)
        writer.writerow(['settings', 'الاختبار النهائي', '', 45, 'minutes', '', 0, 60, 40, '',
                         'single_question', 'rand', '', 200])
        for qa, correct in zip(questions_and_answers, correct_answers):
            question_counter += 1
            qa['answers'] = [i.strip() for i in qa['answers']]
            if shuffle_checker(qa['answers']):
                writer.writerow(['question', qa['question'], '<p><br data-mce-bogus=1></p>', 'single_choice',
                                 2.5, question_counter, 1, 0])
            else:
                writer.writerow(['question', qa['question'], '<p><br data-mce-bogus=1></p>', 'single_choice', 2.5,
                                 question_counter, 1, 1])
            answer_counter = 0
            correct_checker = False
            answers = []
            for j in qa['answers']:
                answer_counter += 1
                if j == correct:
                    answers.append(['answer', j, 'text', 1, 0, '', answer_counter, ''])
                    correct_checker = True
                else:
                    answers.append(['answer', j, 'text', 0, 0, '', answer_counter, ''])
            answer_counter = 0
            if not correct_checker:
                for j in qa['answers']:
                    answer_counter += 1
                    if correct in j:
                        answers[answer_counter - 1][3] = 1
                        correct_checker = True
            writer.writerows(answers)
            if not correct_checker:
                print(f'No answer was found for question {question_counter}')
if style == 2:
    for paragraph in doc.paragraphs:
        line = paragraph.text.strip().split()
        if len(line) == 0:
            continue
        line_text = paragraph.text.strip()
        if paragraph.style.name != 'List Paragraph':
            if current_question:
                questions_and_answers.append({
                    "question": current_question,
                    "answers": current_answers
                })

                # Start a new question
            current_question = line_text
            current_answers = []

            # Otherwise, consider it as an answer and add it to the current_answers list
        elif current_question:
            current_answers.append(line_text)

    acc = 0
    if current_question:
        questions_and_answers.append({
            "question": current_question,
            "answers": current_answers
        })

    counter = 1
    for qa in questions_and_answers:
        input_string = qa['answers']
        if len(input_string) == 2:
            input_string = input_string[0] + '    ' + input_string[1]
        elif len(input_string) == 3:
            input_string = input_string[0] + '    ' + input_string[1] + '     ' + input_string[2]
        elif len(input_string) == 4:
            input_string = input_string[0] + '    ' + input_string[1] + '     ' + input_string[2] \
                           + '      ' + input_string[3]
        elif len(input_string) == 5:
            input_string = input_string[0] + '    ' + input_string[1] + '     ' + input_string[2] \
                           + '      ' + input_string[3] + '      ' + input_string[4]
        split_strings = re.split(r'\s\s', input_string)
        # Filter out any empty strings resulting from consecutive spaces
        split_strings = [s.strip() for s in split_strings if s]
        qa['answers'] = split_strings
        counter += 1
    for qa in questions_and_answers:
        qa['question'] = re.sub('\d+.\s', "", qa['question']).strip()
    question_counter = 0
    with open(output, mode='w', encoding="utf-8-sig", newline='') as file:
        writer = csv.writer(file)
        writer.writerow(['settings', 'الاختبار النهائي', '', 45, 'minutes', '', 0, 60, 40, '',
                         'single_question', 'rand', '', 200])
        for qa in questions_and_answers:
            qa['answers'] = [i.strip() for i in qa['answers']]
            cleaned = [i.replace('*', '').strip() for i in qa['answers']]
            question_counter += 1
            if shuffle_checker(cleaned):
                writer.writerow(['question', qa['question'], '<p><br data-mce-bogus=1></p>', 'single_choice',
                                 2.5, question_counter, 1, 0])
            else:
                writer.writerow(['question', qa['question'], '<p><br data-mce-bogus=1></p>', 'single_choice', 2.5,
                                 question_counter, 1, 1])
            answer_counter = 0
            correct_checker = False
            answers = []
            for j in qa['answers']:
                answer_counter += 1
                if '*' in j:
                    j = j.replace('*', '').strip()
                    answers.append(['answer', j, 'text', 1, 0, '', answer_counter, ''])
                    correct_checker = True
                else:
                    answers.append(['answer', j, 'text', 0, 0, '', answer_counter, ''])
            writer.writerows(answers)
            if not correct_checker:
                print(f'No answer was found for question {question_counter}')

